"""Resume export — convert markdown to PDF and DOCX.

Uses markdown -> HTML -> PDF (via weasyprint) and markdown -> DOCX (via python-docx).
These are optional dependencies — if not installed, export gracefully degrades.
"""

from __future__ import annotations

import re
from pathlib import Path


_PDF_STYLE = """  body { font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 11pt; line-height: 1.35; color: #222; margin: 0; padding: 0; }
  h1 { font-size: 18pt; margin-bottom: 2pt; }
  h2 { font-size: 13pt; margin-top: 8pt; border-bottom: 1px solid #ccc; padding-bottom: 1pt; }
  h3 { font-size: 11pt; margin-top: 4pt; margin-bottom: 2pt; }
  a { color: #222; text-decoration: none; }
  .url { white-space: nowrap; }
  ul { margin-top: 2pt; margin-bottom: 2pt; list-style-position: inside; }
  li { margin-bottom: 1pt; }
  p { margin-top: 3pt; margin-bottom: 3pt; }
  @page { size: Letter; margin: 0.5in; }"""


_BR_RE = re.compile(r"\s*$", re.MULTILINE)

# U+2011 (non-breaking hyphen) — prevents PDF line-wrap splits at the hyphen
_NB_HYPHEN = "\u2011"


def _apply_non_breaking_hyphens(html_body: str, terms: list[str]) -> str:
    """Replace ASCII hyphens with U+2011 in curated compound terms within
    the HTML body, but ONLY in text content — not inside tag attributes,
    URLs, or class names.

    This is a PDF-render-path-only transformation. The markdown source
    and DOCX export are never affected.
    """
    if not terms:
        return html_body

    # Build a case-insensitive regex matching each term as a whole word.
    # Escape regex special chars in terms (none expected, but safe).
    escaped = [re.escape(t) for t in terms]
    pattern = re.compile(
        r"(?<![\w/])(" + "|".join(escaped) + r")(?![\w/])",
        re.IGNORECASE,
    )

    def _replace_in_text(m: re.Match) -> str:
        return m.group(0).replace("-", _NB_HYPHEN)

    # Split HTML into text segments and tag segments, replace only in text
    # Simple approach: split on <...> tags and replace in non-tag parts
    parts = re.split(r"(<[^>]+>)", html_body)
    result = []
    for part in parts:
        if part.startswith("<") and part.endswith(">"):
            # This is a tag — don't modify
            result.append(part)
        else:
            result.append(pattern.sub(_replace_in_text, part))
    return "".join(result)


def _collapse_inline_breaks(md_text: str) -> str:
    """Convert newlines to inline ' · ' separators for company/location/date
    lines and competency bold-label lines, preserving contact block line breaks.

    The contact block (first paragraph after h1, containing email/phone/URLs)
    keeps its line breaks. Everything else that uses single-newline breaks
    within a paragraph gets collapsed to inline separators, saving ~20px per
    eliminated <br> tag.
    """
    lines = md_text.split("\n")
    result: list[str] = []
    in_contact = False
    contact_done = False
    in_competencies = False
    i = 0

    while i < len(lines):
        line = lines[i]

        # Track section headers
        if line.startswith("## "):
            in_competencies = "competenc" in line.lower() or "skill" in line.lower()
            in_contact = False
            contact_done = True
            result.append(line)
            i += 1
            continue

        # Detect contact block: first non-empty content after h1
        if line.startswith("# ") and not contact_done:
            in_contact = False  # h1 itself
            result.append(line)
            i += 1
            # Next non-empty lines until a --- or ## are contact block
            while i < len(lines):
                cl = lines[i]
                if cl.strip() == "---" or cl.startswith("## ") or cl.startswith("# "):
                    break
                if cl.strip():
                    in_contact = True
                result.append(cl)
                i += 1
            contact_done = True
            in_contact = False
            continue

        # Competency bold-label lines: **Label:** skills\n**Label:** skills
        # Collapse continuation lines (multi-line competency entries) into one line
        if in_competencies and line.strip().startswith("**") and ":**" in line:
            # Check if next line is also a competency entry or blank/section
            # This line is a complete competency entry — keep as-is
            result.append(line)
            i += 1
            continue

        # Company/location/date lines under h3 (Professional Experience roles):
        # Pattern: **Company** · Location · *Date*
        # These often span 2 lines with a newline between company and location
        # Only collapse if the current line starts with ** and next line
        # continues with location/date info (no markdown header, no bullet)
        if (
            line.strip().startswith("**")
            and i + 1 < len(lines)
            and lines[i + 1].strip()
            and not lines[i + 1].strip().startswith(("-", "**", "##", "###", "#", "|"))
            and not lines[i + 1].strip() == "---"
            and "·" in (line + " " + lines[i + 1])
        ):
            # Merge with inline separator
            merged = line.rstrip() + " · " + lines[i + 1].strip()
            result.append(merged)
            i += 2
            continue

        result.append(line)
        i += 1

    return "\n".join(result)


def _convert_competencies_table(md_text: str) -> str:
    """Convert markdown tables in competencies/skills sections to labeled plain-text lines.

    If the model emits a pipe table inside a "Core Competencies" or "Skills" section,
    this rewrites it as bold-label paragraphs before HTML conversion, ensuring ATS-safe
    output regardless of model behavior.
    """
    lines = md_text.split('\n')
    result = []
    current_section = ''
    i = 0
    while i < len(lines):
        line = lines[i]

        # Track section headers
        if line.startswith('## '):
            current_section = line[3:].strip().lower()

        # Detect table inside competencies/skills section
        if (
            ('competenc' in current_section or 'skill' in current_section)
            and '|' in line
            and i + 1 < len(lines)
            and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1])
        ):
            # Skip header and separator lines
            i += 2
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if len(cells) >= 2:
                    label = cells[0].strip('*').strip()
                    skills = cells[1]
                    result.append(f'**{label}:** {skills}')
                elif len(cells) == 1 and cells[0].strip():
                    result.append(cells[0])
                i += 1
            continue

        result.append(line)
        i += 1

    return '\n'.join(result)


def export_pdf(
    markdown_path: Path,
    output_path: Path | None = None,
    non_breaking_hyphen_terms: list[str] | None = None,
) -> Path | None:
    """Export a markdown resume to PDF.

    Returns the path to the PDF, or None if export failed (missing deps).

    If non_breaking_hyphen_terms is provided, ASCII hyphens in those
    compound terms are replaced with U+2011 (non-breaking hyphen) in
    the PDF render path ONLY. The markdown source file is never modified.
    """
    try:
        import markdown
        import weasyprint
    except ImportError:
        return None

    if not markdown_path.exists():
        return None

    output_path = output_path or markdown_path.with_suffix(".pdf")
    md_text = markdown_path.read_text()

    # Defense-in-depth: convert any markdown table inside a competencies/skills
    # section to labeled plain-text lines before HTML conversion, so ATS parsers
    # never see a <table> for competencies even if the model emitted pipe syntax.
    md_text = _convert_competencies_table(md_text)
    md_text = _collapse_inline_breaks(md_text)

    # Convert markdown to HTML
    html_body = markdown.markdown(md_text, extensions=["extra", "tables", "nl2br"])

    # Wrap bare URLs in nowrap spans so they don't split across lines
    html_body = re.sub(
        r'(?<!["\'>])(https?://[^\s<]+)',
        r'<span class="url">\1</span>',
        html_body,
    )

    # Apply non-breaking hyphens to curated compound terms (PDF render path only)
    if non_breaking_hyphen_terms:
        html_body = _apply_non_breaking_hyphens(html_body, non_breaking_hyphen_terms)

    # Wrap in a simple HTML document with print-friendly styling
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
{_PDF_STYLE}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    weasyprint.HTML(string=html).write_pdf(str(output_path))
    return output_path


def count_pdf_pages(markdown_text: str) -> int | None:
    """Count the number of PDF pages a markdown resume would produce.

    Uses weasyprint's render() to determine page count without writing a
    file. Returns None if weasyprint or markdown is not installed.

    Applies the same HTML conversion pipeline as export_pdf (competencies
    table conversion, URL wrapping, print-friendly styling) so the page
    count matches what export_pdf would actually produce.
    """
    result = measure_pdf_pages(markdown_text)
    if result is None:
        return None
    return result["page_count"]


def measure_pdf_pages(markdown_text: str, non_breaking_hyphen_terms: list[str] | None = None) -> dict | None:
    """Measure PDF page count and per-page content heights.

    Uses weasyprint's render() to determine the true page layout without
    writing a file. Returns None if weasyprint or markdown is not installed.

    Returns a dict with:
        - page_count: int — number of pages
        - page_heights: list[float] — content height in CSS px per page
        - total_content_height: float — sum of all page content heights
        - printable_page_height: float — printable height of one page in CSS px

    Applies the same HTML conversion pipeline as export_pdf.
    """
    try:
        import markdown
        import weasyprint
    except ImportError:
        return None

    md_text = _convert_competencies_table(markdown_text)
    md_text = _collapse_inline_breaks(md_text)
    html_body = markdown.markdown(md_text, extensions=["extra", "tables", "nl2br"])
    html_body = re.sub(
        r'(?<!["\'>])(https?://[^\s<]+)',
        r'<span class="url">\1</span>',
        html_body,
    )

    # Apply non-breaking hyphens to curated compound terms (PDF render path only)
    if non_breaking_hyphen_terms:
        html_body = _apply_non_breaking_hyphens(html_body, non_breaking_hyphen_terms)

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
{_PDF_STYLE}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    doc = weasyprint.HTML(string=html).render()
    page_count = len(doc.pages)
    page_heights: list[float] = []
    for i, page in enumerate(doc.pages):
        page_box = page._page_box
        if i == page_count - 1:
            # Last page: measure actual content height, not full page box.
            # The body element's height reflects how much content is on
            # this page (may be much less than the full printable area).
            body_height = 0.0
            for child in page_box.children:
                if hasattr(child, "children"):
                    for gc in child.children:
                        if hasattr(gc, "height") and gc.height > body_height:
                            body_height = float(gc.height)
            # If body_height is 0 (couldn't find body), fall back to page_box height
            page_heights.append(body_height if body_height > 0 else float(page_box.height))
        else:
            # Non-final pages: content fills the entire printable area
            page_heights.append(float(page_box.height))

    # Printable page height = the page box height (printable area after margins)
    printable_page_height = float(doc.pages[0]._page_box.height) if doc.pages else 0.0
    total_content_height = sum(page_heights)

    return {
        "page_count": page_count,
        "page_heights": page_heights,
        "total_content_height": total_content_height,
        "printable_page_height": printable_page_height,
    }


def _parse_inline_markdown(text: str):
    """Parse inline markdown (bold, italic, links) into a list of (text, bold, italic) tuples."""
    tokens = []
    pos = 0

    # Pattern order: bold (**...**), italic (*...*), links [text](url)
    pattern = re.compile(
        r'\*\*(.+?)\*\*'  # bold
        r'|\*(.+?)\*'     # italic
        r'|\[(.+?)\]\((.+?)\)',  # link
        re.DOTALL,
    )

    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], False, False, None))

        if m.group(1) is not None:
            # bold
            tokens.append((m.group(1), True, False, None))
        elif m.group(2) is not None:
            # italic
            tokens.append((m.group(2), False, True, None))
        elif m.group(3) is not None:
            # link — display text, with URL
            tokens.append((m.group(3), False, False, m.group(4)))

        pos = m.end()

    if pos < len(text):
        tokens.append((text[pos:], False, False, None))

    return tokens


def _add_runs_to_paragraph(paragraph, text: str, font_name: str = "Arial",
                           font_size: int = 10, color: str = "222222",
                           bold: bool = False, italic: bool = False):
    """Add inline-formatted runs to a paragraph from markdown text."""
    from docx.shared import Pt, RGBColor

    tokens = _parse_inline_markdown(text)
    for token_text, token_bold, token_italic, url in tokens:
        run = paragraph.add_run(token_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold or token_bold
        run.italic = italic or token_italic


def _add_section_rule(paragraph, color: str = "2E5A8C", size: str = "6"):
    """Add a bottom border to a paragraph (section divider line)."""
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)


def _set_cell_shading(cell, color: str):
    """Set cell background shading."""
    from docx.oxml.ns import qn
    from lxml import etree

    tcPr = cell._element.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is None:
        shading = etree.SubElement(tcPr, qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: int = 240):
    """Set tight paragraph spacing (twips). line=240 means single spacing."""
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')


def export_docx(markdown_path: Path, output_path: Path | None = None) -> Path | None:
    """Export a markdown resume to DOCX.

    Returns the path to the DOCX, or None if export failed (missing deps).
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
        from lxml import etree
    except ImportError:
        return None

    if not markdown_path.exists():
        return None

    output_path = output_path or markdown_path.with_suffix(".docx")
    md_text = markdown_path.read_text()
    lines = md_text.split("\n")

    doc = Document()

    # Page setup: US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Default style: Arial 10pt
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor.from_string('222222')
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    ACCENT = '2E5A8C'
    current_section = ''
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip HTML comments
        if line.strip().startswith('<!--'):
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        # Horizontal rule — skip (section spacing handled by heading borders)
        if line.strip() == '---':
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
            # Parse table
            header_cells = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2  # skip header and separator
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                table_rows.append(row_cells)
                i += 1

            # Defense-in-depth: if this table is inside a competencies/skills section,
            # render as labeled plain-text paragraphs instead of a Word table.
            # This ensures ATS-safe output even if the model emits pipe syntax.
            is_competencies = 'competenc' in current_section or 'skill' in current_section

            if is_competencies:
                # Convert each row to a labeled paragraph: **Category:** skills
                for row_data in table_rows:
                    if len(row_data) >= 2:
                        label = row_data[0].strip('**').strip()
                        skills = row_data[1]
                        para = doc.add_paragraph()
                        _set_paragraph_spacing(para, before=0, after=0, line=240)
                        _add_runs_to_paragraph(para, f"**{label}:** {skills}", font_size=10)
                    elif len(row_data) == 1 and row_data[0].strip():
                        para = doc.add_paragraph()
                        _set_paragraph_spacing(para, before=0, after=0, line=240)
                        _add_runs_to_paragraph(para, row_data[0], font_size=10)
                # Small spacer
                spacer = doc.add_paragraph()
                _set_paragraph_spacing(spacer, before=0, after=0, line=240)
                continue

            # Build Word table (non-competencies tables pass through normally)
            num_cols = len(header_cells)

            table = doc.add_table(rows=1 + len(table_rows), cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False

            # Set table indent to 0 so grid aligns with left margin (body text edge)
            tblPr = table._element.find(qn('w:tblPr'))
            if tblPr is not None:
                tblInd = tblPr.find(qn('w:tblInd'))
                if tblInd is None:
                    tblInd = etree.SubElement(tblPr, qn('w:tblInd'))
                tblInd.set(qn('w:w'), '0')
                tblInd.set(qn('w:type'), 'dxa')

                # Zero the table-level left cell margin so cell text starts at the
                # grid edge, aligning with body text at the left margin
                tblCellMar = tblPr.find(qn('w:tblCellMar'))
                if tblCellMar is None:
                    tblCellMar = etree.SubElement(tblPr, qn('w:tblCellMar'))
                left_mar = tblCellMar.find(qn('w:left'))
                if left_mar is None:
                    left_mar = etree.SubElement(tblCellMar, qn('w:left'))
                left_mar.set(qn('w:w'), '0')
                left_mar.set(qn('w:type'), 'dxa')

            # Set column widths
            for col_idx in range(num_cols):
                for row in table.rows:
                    row.cells[col_idx].width = Inches(7.0 / num_cols)

            # Header row
            for col_idx, cell_text in enumerate(header_cells):
                cell = table.rows[0].cells[col_idx]
                _set_cell_shading(cell, 'D5E8F0')
                para = cell.paragraphs[0]
                _set_paragraph_spacing(para, before=20, after=20, line=240)
                _add_runs_to_paragraph(para, cell_text, font_size=9, color=ACCENT, bold=True)

            # Data rows
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= num_cols:
                        break
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    para = cell.paragraphs[0]
                    _set_paragraph_spacing(para, before=20, after=20, line=240)
                    _add_runs_to_paragraph(para, cell_text, font_size=9)

            # Small spacer after table
            spacer = doc.add_paragraph()
            _set_paragraph_spacing(spacer, before=0, after=0, line=240)
            continue

        # H1 — Name (large, centered, accent color)
        if line.startswith('# '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(para, before=0, after=60, line=240)
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor.from_string(ACCENT)
            run.bold = True
            i += 1
            continue

        # H2 — Section header (accent, bottom border)
        if line.startswith('## '):
            text = line[3:].strip()
            current_section = text.lower()
            para = doc.add_paragraph()
            _set_paragraph_spacing(para, before=120, after=40, line=240)
            _add_section_rule(para, color=ACCENT)
            _add_runs_to_paragraph(para, text, font_size=12, color=ACCENT, bold=True)
            i += 1
            continue

        # H3 — Job title (bold, normal color)
        if line.startswith('### '):
            text = line[4:].strip()
            para = doc.add_paragraph()
            _set_paragraph_spacing(para, before=80, after=20, line=240)
            _add_runs_to_paragraph(para, text, font_size=10, color='222222', bold=True)
            i += 1
            continue

        # Bullet point
        if line.startswith('- '):
            text = line[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            _set_paragraph_spacing(para, before=0, after=0, line=240)
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.15)
            _add_runs_to_paragraph(para, text, font_size=10)
            i += 1
            continue

        # Italic-only line (e.g., dates, subtitles)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            text = line.strip('*').strip()
            para = doc.add_paragraph()
            _set_paragraph_spacing(para, before=20, after=20, line=240)
            _add_runs_to_paragraph(para, text, font_size=10, italic=True)
            i += 1
            continue

        # Regular paragraph (may contain bold/italic inline)
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, before=0, after=0, line=240)
        _add_runs_to_paragraph(para, line.strip(), font_size=10)
        i += 1

    doc.save(str(output_path))
    return output_path
